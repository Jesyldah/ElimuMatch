"""Convert REPORT_DRAFT.md to a Word document for Google Docs / Word review."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm, Emu

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "REPORT_DRAFT.md"
OUTPUT = ROOT / "ElimuMatch_Capstone_Report.docx"
OUTPUT_COPY = ROOT / "ElimuMatch_Report.docx"

NAVY = RGBColor(0x1B, 0x3A, 0x4B)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x55, 0x55)
TABLE_HEADER_BG = "264653"


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=BODY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color


def set_paragraph_spacing(p, before=0, after=8, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ElimuMatch  ·  ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run2 = p.add_run()
    run2._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run3 = p.add_run()
    run3._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run4 = p.add_run()
    run4._r.append(fld_end)
    for r in (run2, run3, run4):
        set_run_font(r, size=9, color=MUTED)


INLINE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str, size=11, color=BODY) -> None:
    text = text.replace("\\n", " ")
    parts = INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=size - 1, color=NAVY)
        elif part.startswith("[") and "](" in part:
            label, url = part[1:].rsplit("](", 1)
            url = url.rstrip(")")
            run = paragraph.add_run(f"{label} ({url})" if label != url else url)
            set_run_font(run, size=size, color=RGBColor(0x0B, 0x5C, 0xA8))
            run.underline = True
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_heading(doc, text: str, level: int) -> None:
    """Use real Word Heading styles so Insert/Update TOC works."""
    text = text.strip()
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        set_paragraph_spacing(p, before=18, after=10, line=1.15)
        size = 18
    elif level == 2:
        set_paragraph_spacing(p, before=14, after=6, line=1.15)
        size = 14
    else:
        set_paragraph_spacing(p, before=10, after=4, line=1.15)
        size = 12
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=NAVY)


def add_toc_page(doc) -> None:
    """Insert a Word TOC field after the cover (user updates fields once in Word)."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=12, after=18)
    run = title.add_run("Table of Contents")
    set_run_font(run, size=18, bold=True, color=NAVY)

    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6)

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = paragraph.add_run(
        "Open in Word, then right-click this area → Update Field → Update entire table."
    )
    set_run_font(run, size=10, italic=True, color=MUTED)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    p = doc.add_paragraph()
    br = p.add_run()
    br.add_break(WD_BREAK.PAGE)


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=8, line=1.15)
    add_inline(p, text)


def add_list_item(doc, text: str, numbered=False, number=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    set_paragraph_spacing(p, before=0, after=3, line=1.15)
    prefix = f"{number}. " if numbered else "• "
    run = p.add_run(prefix)
    set_run_font(run, size=11, color=BODY)
    add_inline(p, text)


def add_image(doc, rel_path: str, caption: str | None = None) -> None:
    path = ROOT / rel_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=4)
    if path.exists():
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.3))
    else:
        run = p.add_run(f"[Figure file not found: {rel_path}]")
        set_run_font(run, size=10, italic=True, color=MUTED)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, before=0, after=12)
        add_inline(cap, caption, size=10, color=MUTED)


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=3, after=3, line=1.0)
            val = row[j] if j < len(row) else ""
            add_inline(p, val.strip(), size=9, color=RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else BODY)
            set_cell_borders(cell)
            if i == 0:
                shade_cell(cell, TABLE_HEADER_BG)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:
                shade_cell(cell, "F7F4EF")
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def strip_html(text: str) -> str:
    text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    text = re.sub(r"<a\s+id=\"[^\"]+\"></a>", "", text)
    text = re.sub(r"</?div[^>]*>", "", text)
    text = re.sub(r"<br\s*/?>", "\n", text)
    return text


def preprocess(md: str) -> str:
    md = re.sub(r"<!--.*?-->", "", md, flags=re.S)
    # Drop the draft status banner before the cover
    md = re.sub(
        r"^# ElimuMatch\n## Retention Analytics.*?\n---\n+",
        "",
        md,
        count=1,
        flags=re.S,
    )
    # Drop markdown TOC; Word will use Heading styles + Insert TOC
    md = re.sub(
        r"# Table of Contents\n.*?(?=\n<a id=\"sec-1\"></a>|\n# 1\. Executive Summary)",
        "",
        md,
        count=1,
        flags=re.S,
    )
    md = md.replace("*End of narrative draft.*", "")
    md = md.replace("*End of report.*", "")
    return md


def add_cover(doc, lines: list[str]) -> int:
    """Render the centered cover block; return index after it."""
    # Skip until first H1 after TOC-less cover
    i = 0
    while i < len(lines) and not lines[i].startswith("# ElimuMatch:"):
        i += 1
    cover_bits = []
    while i < len(lines):
        if lines[i].strip() == "---" and cover_bits:
            i += 1
            break
        cover_bits.append(lines[i])
        i += 1

    def spacer(n=1):
        for _ in range(n):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=6)

    spacer(3)
    for line in cover_bits:
        s = line.strip()
        if not s or s.startswith("<") or s == "---":
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if s.startswith("# "):
            set_paragraph_spacing(p, before=12, after=8)
            run = p.add_run(s[2:].strip())
            set_run_font(run, size=26, bold=True, color=NAVY)
        elif s.startswith("## "):
            set_paragraph_spacing(p, before=4, after=18)
            run = p.add_run(s[3:].strip())
            set_run_font(run, size=14, italic=True, color=MUTED)
        elif s.startswith("**") and s.endswith("**"):
            set_paragraph_spacing(p, before=14, after=4)
            run = p.add_run(s.strip("*"))
            set_run_font(run, size=12, bold=True, color=NAVY)
        elif s.startswith("*") and s.endswith("*"):
            set_paragraph_spacing(p, before=2, after=2)
            add_inline(p, s, size=10, color=MUTED)
        else:
            set_paragraph_spacing(p, before=2, after=2)
            add_inline(p, s, size=12)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return i


def convert() -> Path:
    md = preprocess(SOURCE.read_text(encoding="utf-8"))
    lines = [ln.rstrip() for ln in md.splitlines()]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    add_page_number(section)

    i = add_cover(doc, lines)
    add_toc_page(doc)
    in_code = False
    code_buf: list[str] = []
    numbered = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=4, after=10)
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, name="Consolas", size=9, color=NAVY)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("<a id=") or stripped.startswith("<div") or stripped.startswith("</div>"):
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if not stripped:
            numbered = 0
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        m_img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m_img:
            add_image(doc, m_img.group(2), caption=m_img.group(1) or None)
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            numbered = int(stripped.split(".", 1)[0])
            add_list_item(doc, stripped.split(".", 1)[1].strip(), numbered=True, number=numbered)
            i += 1
            continue
        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:])
            i += 1
            continue

        add_body(doc, stripped)
        i += 1

    try:
        doc.save(OUTPUT)
        primary = OUTPUT
    except OSError:
        primary = ROOT / "ElimuMatch_Investor_Brief.docx"
        doc.save(primary)
        print(f"Note: {OUTPUT.name} is locked; wrote {primary.name} instead.")

    for dest in (OUTPUT_COPY, ROOT / "ElimuMatch_Investor_Brief.docx"):
        if dest.resolve() == primary.resolve():
            continue
        try:
            import shutil

            shutil.copy2(primary, dest)
        except OSError as exc:
            print(f"Warning: could not copy to {dest.name}: {exc}")
    return primary


if __name__ == "__main__":
    out = convert()
    print(f"Wrote {out.name} ({out.stat().st_size / 1024:.0f} KB)")
    for name in ("ElimuMatch_Report.docx", "ElimuMatch_Investor_Brief.docx", "ElimuMatch_Capstone_Report.docx"):
        p = ROOT / name
        if p.exists() and p.resolve() != out.resolve():
            print(f"Also available: {name}")
